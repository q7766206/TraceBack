"""
TraceBack Report API 路由
提供回溯分析报告的生成、状态查询和获取接口
"""

import os
import json
import uuid
import threading
import time
from datetime import datetime
from flask import request, jsonify

from . import report_bp
from ..config import Config
from ..utils.llm_client import LLMClient
from ..utils.logger import get_logger

logger = get_logger('traceback.api.report')

# 报告存储
_reports = {}  # report_id -> report data
_report_tasks = {}  # report_id -> generation status


def _get_report_dir(report_id):
    d = os.path.join(Config.UPLOAD_FOLDER, 'reports', report_id)
    os.makedirs(d, exist_ok=True)
    return d


def _save_report(report_id, data):
    _reports[report_id] = data
    d = _get_report_dir(report_id)
    with open(os.path.join(d, 'report.json'), 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _load_report(report_id):
    if report_id in _reports:
        return _reports[report_id]
    d = _get_report_dir(report_id)
    fp = os.path.join(d, 'report.json')
    if os.path.exists(fp):
        with open(fp, 'r', encoding='utf-8') as f:
            data = json.load(f)
            _reports[report_id] = data
            return data
    return None


def _load_simulation(simulation_id):
    """加载 simulation 数据"""
    sim_dir = os.path.join(Config.UPLOAD_FOLDER, 'simulations', simulation_id)
    fp = os.path.join(sim_dir, 'simulation.json')
    if os.path.exists(fp):
        with open(fp, 'r', encoding='utf-8') as f:
            return json.load(f)
    return None


def _load_run_result(simulation_id):
    """加载模拟运行结果"""
    sim_dir = os.path.join(Config.UPLOAD_FOLDER, 'simulations', simulation_id)
    fp = os.path.join(sim_dir, 'run_result.json')
    if os.path.exists(fp):
        with open(fp, 'r', encoding='utf-8') as f:
            return json.load(f)
    # 也尝试从引擎内存获取
    from ..services.simulation_engine import get_run_state
    state = get_run_state(simulation_id)
    if state and state.all_actions:
        return {
            "actions": state.all_actions,
            "posts": state.posts,
            "total_actions": len(state.all_actions),
        }
    return None


@report_bp.route('/generate', methods=['POST'])
def generate_report():
    """生成回溯分析报告"""
    try:
        data = request.get_json()
        simulation_id = data.get('simulation_id')
        force = data.get('force_regenerate', False)

        if not simulation_id:
            return jsonify({"success": False, "error": "缺少 simulation_id"}), 400

        # 检查是否已有报告
        if not force:
            for rid, rdata in _reports.items():
                if rdata.get('simulation_id') == simulation_id and rdata.get('status') in ('generating', 'completed'):
                    return jsonify({"success": True, "data": {"report_id": rid, "status": rdata['status']}})

        report_id = f"rpt_{uuid.uuid4().hex[:12]}"

        report_data = {
            "report_id": report_id,
            "simulation_id": simulation_id,
            "status": "generating",
            "created_at": datetime.now().isoformat(),
            "outline": None,
            "sections": {},
            "current_section": 0,
            "total_sections": 0,
            "agent_logs": [],
            "console_logs": [],
        }
        _save_report(report_id, report_data)

        # 启动后台生成
        thread = threading.Thread(
            target=_generate_report_worker,
            args=(report_id, simulation_id),
            daemon=True,
        )
        thread.start()

        logger.info(f"报告生成启动: {report_id} (sim={simulation_id})")
        return jsonify({"success": True, "data": {"report_id": report_id, "status": "generating"}})

    except Exception as e:
        logger.error(f"报告生成失败: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


@report_bp.route('/generate/status', methods=['GET'])
def get_report_status():
    """获取报告生成状态"""
    report_id = request.args.get('report_id')
    if not report_id:
        return jsonify({"success": False, "error": "缺少 report_id"}), 400

    report = _load_report(report_id)
    if not report:
        return jsonify({"success": False, "error": "报告不存在"}), 404

    return jsonify({
        "success": True,
        "data": {
            "report_id": report_id,
            "status": report.get("status", "unknown"),
            "current_section": report.get("current_section", 0),
            "total_sections": report.get("total_sections", 0),
            "outline": report.get("outline"),
        }
    })


@report_bp.route('/<report_id>', methods=['GET'])
def get_report(report_id: str):
    """获取完整报告"""
    report = _load_report(report_id)
    if not report:
        return jsonify({"success": False, "error": "报告不存在"}), 404

    return jsonify({"success": True, "data": report})


@report_bp.route('/<report_id>/agent-log', methods=['GET'])
def get_agent_log(report_id: str):
    """获取 Agent 日志（增量）"""
    from_line = int(request.args.get('from_line', 0))
    report = _load_report(report_id)
    if not report:
        return jsonify({"success": False, "error": "报告不存在"}), 404

    logs = report.get("agent_logs", [])
    new_logs = logs[from_line:]
    return jsonify({"success": True, "data": {"logs": new_logs, "total": len(logs), "from_line": 0}})


@report_bp.route('/<report_id>/console-log', methods=['GET'])
def get_console_log(report_id: str):
    """获取控制台日志（增量）"""
    from_line = int(request.args.get('from_line', 0))
    report = _load_report(report_id)
    if not report:
        return jsonify({"success": False, "error": "报告不存在"}), 404

    logs = report.get("console_logs", [])
    new_logs = logs[from_line:]
    return jsonify({"success": True, "data": {"logs": new_logs, "total": len(logs), "from_line": 0}})





@report_bp.route('/<report_id>/download', methods=['GET'])
def download_report(report_id: str):
    """下载报告为 docx 文档"""
    import io
    import re
    from urllib.parse import quote
    from flask import Response, send_file
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    report = _load_report(report_id)
    if not report:
        return jsonify({"success": False, "error": "报告不存在"}), 404

    outline = report.get("outline")
    if not outline:
        return jsonify({"success": False, "error": "报告尚未生成完成"}), 400

    # 从 agent_logs 中提取各章节内容
    sections_content = {}
    agent_logs = report.get("agent_logs", [])
    for log in agent_logs:
        if log.get("action") == "section_complete" and log.get("details", {}).get("content"):
            sec_idx = log.get("section_index")
            if sec_idx is not None:
                raw = log["details"]["content"]
                sections_content[sec_idx] = raw if isinstance(raw, str) else (raw.get("content") if isinstance(raw, dict) else str(raw))

    title = outline.get("title", "TraceBack Report")
    summary = outline.get("summary", "")
    sections = outline.get("sections", [])

    # ── 构建 docx ──
    doc = Document()

    # 页面设置
    section_page = doc.sections[0]
    section_page.top_margin = Cm(2.54)
    section_page.bottom_margin = Cm(2.54)
    section_page.left_margin = Cm(3.18)
    section_page.right_margin = Cm(3.18)

    # 样式配置
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)

    # 封面标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(72)
    p_title.paragraph_format.space_after = Pt(24)
    run_title = p_title.add_run(title)
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # 摘要
    if summary:
        p_summary = doc.add_paragraph()
        p_summary.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_summary.paragraph_format.space_after = Pt(12)
        p_summary.paragraph_format.left_indent = Cm(1)
        p_summary.paragraph_format.right_indent = Cm(1)
        run_sum = p_summary.add_run(summary)
        run_sum.font.size = Pt(11)
        run_sum.font.italic = True
        run_sum.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    # 元信息
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(24)
    meta_text = f"Report ID: {report_id}    |    Generated: {report.get('created_at', 'N/A')[:19]}    |    Status: {report.get('status', 'unknown')}"
    run_meta = p_meta.add_run(meta_text)
    run_meta.font.size = Pt(9)
    run_meta.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # 分隔线
    doc.add_paragraph('─' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Markdown 转 docx 段落的辅助函数 ──
    def _add_md_content(doc, md_text):
        """将 Markdown 文本转为 docx 段落"""
        if not md_text:
            return

        lines = md_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # 跳过空行
            if not stripped:
                i += 1
                continue

            # 标题
            if stripped.startswith('#### '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(stripped[5:].strip())
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
                i += 1
                continue
            if stripped.startswith('### '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(stripped[4:].strip())
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
                i += 1
                continue
            if stripped.startswith('## '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(stripped[3:].strip())
                run.font.size = Pt(15)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
                i += 1
                continue

            # 引用块
            if stripped.startswith('> '):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.5)
                p.paragraph_format.space_after = Pt(8)
                quote_text = stripped[2:].strip()
                # 处理粗体
                _add_rich_runs(p, quote_text, italic=True, color=RGBColor(0x4B, 0x55, 0x63))
                i += 1
                continue

            # 无序列表
            if stripped.startswith('- ') or stripped.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(3)
                _add_rich_runs(p, stripped[2:].strip())
                i += 1
                continue

            # 有序列表
            ol_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
            if ol_match:
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_after = Pt(3)
                _add_rich_runs(p, ol_match.group(2).strip())
                i += 1
                continue

            # 普通段落
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _add_rich_runs(p, line)
            i += 1


    reports_dir = os.path.join(Config.UPLOAD_FOLDER, 'reports')
    if os.path.exists(reports_dir):
        for d in os.listdir(reports_dir):
            fp = os.path.join(reports_dir, d, 'report.json')
            if os.path.exists(fp):
                try:
                    with open(fp, 'r', encoding='utf-8') as f:
                        r = json.load(f)
                    reports.append({
                        "report_id": r.get("report_id"),
                        "simulation_id": r.get("simulation_id"),
                        "status": r.get("status"),
                        "created_at": r.get("created_at"),
                    })
                except:
                    pass
    return jsonify({"success": True, "data": reports, "count": len(reports)})


# ═══════════════════════════════════════════════════════════════
# 后台报告生成 Worker
# ═══════════════════════════════════════════════════════════════

def _add_agent_log(report_id, action, section_index=None, details=None, message=""):
    """添加结构化 Agent 日志（前端 Step4Report 期望的格式）"""
    report = _load_report(report_id)
    if not report:
        return
    entry = {
        "timestamp": datetime.now().isoformat(),
        "action": action,
        "message": message,
    }
    if section_index is not None:
        entry["section_index"] = section_index
    if details is not None:
        entry["details"] = details
    report.setdefault("agent_logs", []).append(entry)
    _save_report(report_id, report)


def _add_console_log(report_id, message):
    """添加控制台日志"""
    report = _load_report(report_id)
    if not report:
        return
    report.setdefault("console_logs", []).append(message)
    _save_report(report_id, report)


def _generate_report_worker(report_id, simulation_id):
    """后台生成报告"""
    try:
        llm = LLMClient()
        report = _load_report(report_id)

        _add_console_log(report_id, "开始生成回溯分析报告...")
        _add_agent_log(report_id, "report_start", message="Report Agent 已启动")

        # 1. 加载模拟数据
        sim = _load_simulation(simulation_id)
        run_result = _load_run_result(simulation_id)

        requirement = ""
        if sim:
            requirement = sim.get("simulation_requirement", "")

        actions_summary = ""
        if run_result:
            actions = run_result.get("actions", [])
            _add_console_log(report_id, f"加载到 {len(actions)} 条模拟动作")

            posts = [a for a in actions if a.get("action_type") == "CREATE_POST" and a.get("content")]
            for p in posts[:20]:
                actions_summary += f"\n[@{p.get('agent_name','?')}|{p.get('platform','?')}|R{p.get('round_num',0)}] {p.get('content','')[:150]}"
        else:
            _add_console_log(report_id, "警告：未找到模拟运行结果，使用有限数据生成报告")

        # 2. 生成报告大纲
        _add_agent_log(report_id, "planning_start", message="正在生成报告大纲...")
        _add_console_log(report_id, "正在调用 LLM 生成大纲...")

        outline_prompt = f"""你是一个专业的历史事件因果分析报告撰写专家。

研究课题: {requirement}

模拟讨论摘要（来自多 Agent 社交媒体模拟）:
{actions_summary[:3000]}

请生成一份回溯分析报告的大纲，以 JSON 格式输出：
{{
    "title": "报告标题",
    "summary": "200字以内的报告摘要",
    "sections": [
        {{"title": "章节标题", "description": "章节描述", "key_points": ["要点1", "要点2"]}}
    ]
}}

要求：
1. 包含 4-6 个章节
2. 覆盖：事件概述、关键时间线、因果分析、多方观点、争议焦点、结论与建议
3. 基于模拟讨论中的实际内容"""

        outline = llm.chat_json([{"role": "user", "content": outline_prompt}])
        if isinstance(outline, list):
            outline = outline[0] if outline else {}

        report = _load_report(report_id)
        report["outline"] = outline
        report["total_sections"] = len(outline.get("sections", []))
        _save_report(report_id, report)

        # 发送 planning_complete 事件（前端据此渲染大纲）
        _add_agent_log(report_id, "planning_complete", details={"outline": outline},
                       message=f"大纲生成完成: {outline.get('title', '?')}")
        _add_console_log(report_id, f"报告共 {report['total_sections']} 个章节")

        # 3. 逐章节生成内容
        sections = outline.get("sections", [])
        previous_sections_text = []  # 已完成章节的文本，供后续章节参考避免重复
        for idx, section in enumerate(sections):
            section_num = idx + 1
            report = _load_report(report_id)
            report["current_section"] = section_num
            _save_report(report_id, report)

            # section_start 事件
            _add_agent_log(report_id, "section_start", section_index=section_num,
                           message=f"正在撰写第 {section_num} 章: {section.get('title', '?')}")
            _add_console_log(report_id, f"开始生成第 {section_num}/{len(sections)} 章...")

            # 构建已完成章节摘要（每章最多800字），帮助LLM避免重复
            prev_summary = ""
            if previous_sections_text:
                prev_parts = []
                for ps in previous_sections_text:
                    prev_parts.append(ps[:800] + "..." if len(ps) > 800 else ps)
                prev_summary = "\n\n---\n\n".join(prev_parts)

            section_prompt = f"""你是历史事件因果分析报告撰写专家。请撰写以下章节的完整内容。

报告标题: {outline.get('title', '')}
当前章节: 第 {section_num} 章 - {section.get('title', '')}
章节描述: {section.get('description', '')}
关键要点: {json.dumps(section.get('key_points', []), ensure_ascii=False)}

研究背景: {requirement[:500]}

模拟讨论数据:
{actions_summary[:2000]}

{"已完成的章节内容（请避免重复）:" + chr(10) + prev_summary if prev_summary else "（这是第一个章节）"}

请直接撰写章节正文内容（800-1500字），使用 Markdown 格式：
- 用段落形式组织内容，不要用纯列表堆砌
- 用 **粗体** 标记关键概念
- 用 > 引用格式展示重要引述
- 不要添加章节标题（系统会自动添加）
- 不要输出 JSON，直接输出正文"""

            section_content = None
            max_retries = 2
            for retry in range(max_retries + 1):
                try:
                    raw_content = llm.chat(
                        [{"role": "user", "content": section_prompt}],
                        temperature=0.7,
                        max_tokens=4096,
                    )
                    if raw_content and len(raw_content.strip()) > 50:
                        section_content = {
                            "title": section.get("title", ""),
                            "content": raw_content.strip(),
                            "key_findings": [],
                            "evidence": [],
                        }
                        break
                    else:
                        logger.warning(f"章节 {section_num} 第{retry+1}次生成内容过短({len(raw_content.strip()) if raw_content else 0}字)，重试...")
                except Exception as e:
                    logger.warning(f"章节 {section_num} 第{retry+1}次生成失败: {e}")
                    if retry < max_retries:
                        time.sleep(3)

            if not section_content:
                logger.error(f"章节 {section_num} 多次重试后仍失败，写入降级内容")
                section_content = {
                    "title": section.get("title", ""),
                    "content": f"本章节内容生成失败，请稍后重试。章节主题：{section.get('description', '')}",
                    "key_findings": [], "evidence": [],
                }

            # 保存章节（无论成功或降级，都推进进度）
            report = _load_report(report_id)
            report.setdefault("sections", {})[str(section_num)] = section_content
            _save_report(report_id, report)

            content_len = len(section_content.get('content', ''))
            _add_agent_log(report_id, "section_complete", section_index=section_num,
                           details={"content": section_content},
                           message=f"第 {section_num} 章撰写完成")
            _add_console_log(report_id, f"第 {section_num} 章完成 ({content_len}字)")

            # 记录已完成章节文本
            previous_sections_text.append(f"## {section_content['title']}\n\n{section_content['content']}")

        # 4. 完成
        report = _load_report(report_id)
        report["status"] = "completed"
        report["completed_at"] = datetime.now().isoformat()
        _save_report(report_id, report)

        _add_agent_log(report_id, "report_complete", message="报告生成完成")
        _add_console_log(report_id, f"报告已完成: {report_id}")
        logger.info(f"报告生成完成: {report_id}")

    except Exception as e:
        logger.error(f"报告生成失败: {report_id}: {e}", exc_info=True)
        report = _load_report(report_id)
        if report:
            report["status"] = "failed"
            report["error"] = str(e)
            _save_report(report_id, report)
        _add_console_log(report_id, f"报告生成失败: {str(e)}")
    def _add_rich_runs(paragraph, text, italic=False, color=None):
        """处理 Markdown 粗体/斜体标记，生成多个 run"""
        # 匹配 **bold** 和 *italic*
        pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*)')
        last_end = 0
        for m in pattern.finditer(text):
            # 前面的普通文本
            if m.start() > last_end:
                run = paragraph.add_run(text[last_end:m.start()])
                run.font.size = Pt(12)
                if italic:
                    run.font.italic = True
                if color:
                    run.font.color.rgb = color
            # 粗体
            if m.group(2):
                run = paragraph.add_run(m.group(2))
                run.font.size = Pt(12)
                run.font.bold = True
                if italic:
                    run.font.italic = True
                if color:
                    run.font.color.rgb = color
            # 斜体
            elif m.group(3):
                run = paragraph.add_run(m.group(3))
                run.font.size = Pt(12)
                run.font.italic = True
                if color:
                    run.font.color.rgb = color
            last_end = m.end()
        # 剩余文本
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.size = Pt(12)
            if italic:
                run.font.italic = True
            if color:
                run.font.color.rgb = color

    # ── 各章节 ──
    for idx, sec in enumerate(sections):
        section_num = idx + 1
        sec_title = sec.get("title", f"Section {section_num}")

        # 章节标题
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(24)
        p_sec.paragraph_format.space_after = Pt(12)
        run_sec = p_sec.add_run(sec_title)
        run_sec.font.size = Pt(18)
        run_sec.font.bold = True
        run_sec.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

        content = sections_content.get(section_num)
        if content:
            _add_md_content(doc, content)
        else:
            p_empty = doc.add_paragraph()
            run_empty = p_empty.add_run("（本章节内容尚未生成）")
            run_empty.font.italic = True
            run_empty.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # 输出到内存
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = re.sub(r'[\\/:*?"<>|]', '-', title[:50]).strip()
    filename = f"{safe_title}_{report_id}.docx"

    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename,
    )


@report_bp.route('/list', methods=['GET'])
def list_reports():
    """列出报告"""
    reports = []
    reports_dir = os.path.join(Config.UPLOAD_FOLDER, 'reports')
    if os.path.exists(reports_dir):
        for d in os.listdir(reports_dir):
            fp = os.path.join(reports_dir, d, 'report.json')
            if os.path.exists(fp):
                try:
                    with open(fp, 'r', encoding='utf-8') as f:
                        r = json.load(f)
                    reports.append({
                        "report_id": r.get("report_id"),
                        "simulation_id": r.get("simulation_id"),
                        "status": r.get("status"),
                        "created_at": r.get("created_at"),
                    })
                except:
                    pass
    return jsonify({"success": True, "data": reports, "count": len(reports)})


# ═══════════════════════════════════════════════════════════════
# 后台报告生成 Worker
# ═══════════════════════════════════════════════════════════════
